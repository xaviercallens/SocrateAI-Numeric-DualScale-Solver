#!/usr/bin/env python3
"""
Génération des documents DOCX conformes aux spécifications de l'INPI
(Institut National de la Propriété Industrielle) pour la demande de brevet :
"Procédé, système informatique et dispositif matériel pour la simulation numérique
et le contrôle prédictif en temps réel d'écoulements fluides par régularisation spectrale à double échelle."

Résolution intégrale des anomalies INPI :
1. Aucune formule/image manquante pour [Math 1], [Math 2], [Math 3] :
   Génération et insertion d'images PNG haute résolution (300 DPI) au paragraphe suivant chaque balise.
2. Élimination du symbole U+207B (exposant moins) :
   Remplacement de 10⁻¹⁶ et 10⁻¹⁴ par la notation standard ASCII 10^-16 et 10^-14.
3. Rétablissement de la numérotation stricte des paragraphes :
   Suppression du paragraphe de métadonnées hors-style (inventeur/date) au début de la description,
   garantissant une suite ininterrompue [0001] à [0024] conforme au PDF généré par l'INPI.
"""

import os
import shutil
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt

def ensure_math_images(fig_dir):
    """Génère les images PNG nettes (300 DPI) des formules mathématiques balisées."""
    os.makedirs(fig_dir, exist_ok=True)
    
    math_defs = [
        ("math1.png", r"D(k) = -\nu |k|^2 \cdot \left[1 + \alpha^\prime |k|^2\right]"),
        ("math2.png", r"P(k) = I - \frac{k \otimes k}{|k|^2}"),
        ("math3.png", r"I_{\mathrm{frust}} = \frac{\sum |T(k,p,q)|}{\left|\sum T(k,p,q)\right|}")
    ]
    
    for filename, latex_str in math_defs:
        filepath = os.path.join(fig_dir, filename)
        fig = plt.figure(figsize=(5.5, 0.9), dpi=300)
        plt.text(0.5, 0.5, f"${latex_str}$", fontsize=15, ha="center", va="center")
        plt.axis("off")
        plt.savefig(filepath, bbox_inches="tight", transparent=False, facecolor="white")
        plt.close()
    print("Images des formules mathématiques [Math 1], [Math 2], [Math 3] prêtes.")

def clear_body(doc):
    """Vide le corps du document tout en préservant w:sectPr (mise en page, marges)."""
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("p") or child.tag.endswith("tbl"):
            body.remove(child)

def add_num_pr(paragraph, num_id=2, ilvl=0):
    """Ajoute la balise w:numPr pour la numérotation automatique INPI."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_elem = OxmlElement('w:ilvl')
    ilvl_elem.set(qn('w:val'), str(ilvl))
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    pPr.append(numPr)

def add_subtitle(doc, text, style):
    """Ajoute un sous-titre conforme INPI (< 50 caractères, gras, style sous-titre)."""
    assert len(text) < 50, f"Sous-titre INPI trop long ({len(text)} car >= 50) : {text}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = True
    return p

def generate_libreoffice_docx(template_path, output_path, math_dir, img_fig1=None, img_fig2=None):
    """Génère le document DOCX INPI au format LibreOffice (Open Source)."""
    doc = docx.Document(template_path)
    clear_body(doc)

    # 1. SECTION DESCRIPTION
    doc.add_paragraph("Description", style="Inpi_Section")
    
    # Titre de l'invention (strictement suivi du premier sous-titre)
    titre_text = ("Titre de l’invention : Procédé, système informatique et dispositif matériel "
                  "pour la simulation numérique et le contrôle prédictif en temps réel d'écoulements "
                  "fluides par régularisation spectrale à double échelle.")
    doc.add_paragraph(titre_text, style="Inpi_Titre")

    # Domaine technique
    add_subtitle(doc, "Domaine technique de l'invention", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "La présente invention se rapporte au domaine de la mécanique des fluides numérique "
        "(CFD - Computational Fluid Dynamics), des systèmes cyber-physiques et de l'informatique embarquée. "
        "Plus particulièrement, l'invention concerne un procédé mis en œuvre par ordinateur pour la résolution "
        "stabilisée des équations de Navier-Stokes, optimisé pour s'exécuter sur des architectures matérielles "
        "à ressources limitées (microcontrôleurs « bare-metal ») afin de piloter et d'asservir des équipements "
        "physiques fluides en temps réel (ex : bioréacteurs, aérodynamique active, refroidissement thermique).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # État de la technique
    add_subtitle(doc, "État de la technique", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "La simulation d'écoulements turbulents pour le contrôle industriel repose sur les équations de Navier-Stokes. "
        "Le problème technique majeur réside dans la formation de structures tourbillonnaires à des échelles de plus en "
        "plus petites (phénomène d'étirement tourbillonnaire ou vortex stretching), ce qui conduit fréquemment dans les "
        "solveurs numériques à des singularités de calcul, des instabilités par accumulation d'énergie aux petites échelles, "
        "des dépassements de capacité produisant des erreurs fatales de type NaN (Not a Number), et par conséquent au plantage "
        "des systèmes informatiques de contrôle industriel.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Pour tenter de contourner ces limitations, les logiciels de l'état de l'art (modèles LES, solveurs aux Volumes Finis "
        "tels qu'OpenFOAM ou ANSYS Fluent) nécessitent d'une part des résolutions itératives de systèmes matriciels creux "
        "(équation de Poisson pour la pression) exigeant plusieurs Gigaoctets de mémoire vive (RAM), ce qui interdit formellement "
        "tout déploiement embarqué sur microcontrôleurs pour un asservissement à faible latence (Edge Computing). D'autre part, "
        "ils recourent à des viscosités artificielles ou des filtres empiriques de sous-maille qui ne garantissent mathématiquement "
        "aucun bornage strict de l'enstrophie, laissant subsister un risque majeur d'interruption du processeur lors de fluctuations critiques.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Il existe donc un besoin technique impérieux pour un procédé de calcul et de contrôle fluide garantissant informatiquement "
        "l'absence totale de dépassement de capacité (stabilité inconditionnelle et bornage de l'enstrophie), tout en réduisant drastiquement "
        "l'empreinte mémoire à quelques kilo-octets pour permettre un contrôle matériel en boucle fermée ultra-rapide.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Exposé de l'invention
    add_subtitle(doc, "Exposé de l'invention", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "L'invention résout ce problème technique en introduisant un moteur de calcul pseudo-spectral couplé à une régularisation "
        "ultraviolette à plancher d'échelle (fondée sur une topologie de double échelle), implémenté sous la forme d'un code système "
        "déterministe sans allocation dynamique de mémoire (fonctionnant en mode no_std).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Description détaillée
    add_subtitle(doc, "Description détaillée", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "Le procédé selon l'invention comprend les étapes techniques suivantes, exécutées en boucle temps réel par au moins un processeur :",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Étape A : Acquisition et Transformation spectrale. Le système acquiert les signaux de capteurs physiques (vitesses, pressions "
        "locales, consignes) et transforme le champ d'écoulement dans un espace spectral de Fourier à l'aide d'une transformée rapide (FFT).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Étape B : Filtrage par Dissipation Modifiée (Plancher d'échelle). Le processeur applique itérativement un opérateur de dissipation "
        "spectrale D(k) configuré selon une loi modifiée à régularisation biharmonique d'ordre 4 aux hautes fréquences :",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Balise mathématique 1 + Image de la formule
    doc.add_paragraph("[Math 1]", style="Inpi_BalisesConditionnelles")
    p_m1 = doc.add_paragraph(style="Inpi_ParagraphesNonNumérotés")
    p_m1.add_run().add_picture(os.path.join(math_dir, "math1.png"), height=Inches(0.4))

    p = doc.add_paragraph(
        "où ν désigne la viscosité cinématique moléculaire, k le vecteur d'onde spatial, et α' un paramètre dimensionnel fixant un plancher "
        "d'échelle ultraviolet. L'effet technique direct de cette étape est la garantie informatique du maintien de l'enstrophie sous une borne "
        "stricte indépendante du maillage, interdisant mathématiquement et matériellement tout dépassement de capacité des registres en virgule flottante "
        "et tout plantage du processeur.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Étape C : Intégration temporelle exacte et Projection de Leray-Helmholtz. Le processeur intègre temporellement l'état du fluide "
        "par un intégrateur exponentiel analytique (schéma ETD-RK4 ou IF-RK2 de Lawson) qui traite exactement la partie linéaire visqueuse, "
        "affranchissant le calcul de la condition de stabilité de Courant-Friedrichs-Lewy (CFL) visqueuse. À chaque pas, une projection algébrique "
        "orthogonale de Leray-Helmholtz est appliquée directement dans le domaine de Fourier selon :",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Balise mathématique 2 + Image de la formule
    doc.add_paragraph("[Math 2]", style="Inpi_BalisesConditionnelles")
    p_m2 = doc.add_paragraph(style="Inpi_ParagraphesNonNumérotés")
    p_m2.add_run().add_picture(os.path.join(math_dir, "math2.png"), height=Inches(0.4))

    p = doc.add_paragraph(
        "L'effet technique de cette projection algébrique directe est de garantir une divergence strictement nulle (incompressibilité à l'epsilon "
        "machine près, inférieure à 10^-16) pour chaque mode de Fourier, sans nécessiter le moindre stockage matriciel creux ni la moindre "
        "itération de solveur de Poisson en mémoire vive (supprimant les algorithmes itératifs de type SIMPLE ou PISO).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Étape D : Capteur logiciel de Frustration Triadique. Le processeur calcule dynamiquement, à partir des transferts non linéaires d'énergie, "
        "un indice adimensionnel dit « Indice de Frustration Triadique », défini par le rapport :",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Balise mathématique 3 + Image de la formule
    doc.add_paragraph("[Math 3]", style="Inpi_BalisesConditionnelles")
    p_m3 = doc.add_paragraph(style="Inpi_ParagraphesNonNumérotés")
    p_m3.add_run().add_picture(os.path.join(math_dir, "math3.png"), height=Inches(0.4))

    p = doc.add_paragraph(
        "Cet indice mesure le degré de désalignement géométrique des triades de turbulence. Lorsque cet indice franchit un seuil prédéterminé, "
        "le système déclenche une allocation dynamique des ressources matérielles de calcul ou adapte la fréquence d'échantillonnage temporel.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "Étape E : Asservissement matériel en boucle fermée. L'état prédit de l'écoulement est converti en temps réel en un signal électrique de commande "
        "transmis à un actionneur physique (e.g. variateur de vitesse d'agitateur de bioréacteur pour asservir le transfert massique kLa, volets "
        "d'aérodynamique active, pompes de refroidissement thermique).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # Avantages techniques
    add_subtitle(doc, "Avantages techniques", "INPI_Balises_personnalisées")
    p15_text = (
        "La synergie entre la régularisation spectrale biharmonique, la projection algébrique de Leray-Helmholtz et l'architecture sans allocation dynamique "
        "permet des performances matérielles inédites mesurées sur banc d'essai : "
        "(1) Une empreinte de mémoire vive (RAM) inférieure à 3 Kilo-octets (mesurée précisément à 2 624 octets pour une grille de dimensionnement embarqué), "
        "permettant une implantation directe dans la mémoire SRAM interne de microcontrôleurs basse consommation (architectures ARM Cortex-M4/M7 ou RISC-V), "
        "sans mémoire externe ; "
        "(2) Un déterminisme d'exécution temps réel strict avec une latence médiane par itération inférieure à 100 microsecondes (mesurée à 59,8 µs), "
        "autorisant des boucles de contrôle cyber-physiques à plus de 10 000 Hz."
    )
    p = doc.add_paragraph(p15_text, style="Inpi_ParagraphesNumérotés")
    add_num_pr(p, num_id=2, ilvl=0)

    # Optimisation autonome et benchmarks
    add_subtitle(doc, "Validation empirique certifiée", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "Le procédé s'intègre en outre dans un système d'optimisation par recherche autonome et modèles d'ordre réduit (ROMs), "
        "où chaque configuration paramétrique est validée vis-à-vis d'invariants formels certifiés par démonstrateur interactif de théorèmes (Lean 4). "
        "L'ensemble des résultats empiriques a été scellé par une empreinte cryptographique SHA-256 (Sceau : 0ae0e5d97da424e0).",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p17_text = (
        "Les gains quantitatifs établis par rapport aux solveurs conventionnels (Volumes Finis type OpenFOAM) comprennent : "
        "(i) Précision numérique absolue (Zéro diffusion numérique) : Sur le cas test standardisé du Tourbillon de Taylor-Green (UC7), "
        "l'erreur L2 sur la vitesse est mesurée à 7,24 x 10^-14 sur grille 128², prouvant l'éradication totale de la viscosité artificielle de discrétisation ; "
        "(ii) Vitesse algorithmique (Contournement de la limite CFL) : Résolution de la cavité entraînée (UC8) en 0,59 seconde et de la cascade "
        "de Kolmogorov en 3D (UC11) en 0,21 seconde grâce au schéma intégrateur exponentiel ETD-RK4, soit un gain de vitesse de calcul supérieur à 10× "
        "par rapport aux algorithmes itératifs de pression ; "
        "(iii) Sécurité épistémique matérielle (Pare-feu mathématique anti-hallucination) : Soumis à des conditions aux limites non physiques "
        "(e.g. UC15 - tore périodique avec vorticité moyenne non nulle forçant un paradoxe mathématique), le procédé rejette formellement l'état "
        "en mesurant une perte de circulation de ~100% et déclenche une interruption de sécurité matérielle « Échec Contrôle Négatif », "
        "empêchant toute prise de commande erronée par l'actionneur physique."
    )
    p = doc.add_paragraph(p17_text, style="Inpi_ParagraphesNumérotés")
    add_num_pr(p, num_id=2, ilvl=0)

    # Brève description des dessins
    add_subtitle(doc, "Brève description des dessins", "INPI_Balises_personnalisées")
    p = doc.add_paragraph(
        "[Fig 1] représente un graphique comparatif de la décroissance temporelle de l'énergie cinétique entre le procédé selon l'invention "
        "(LeanFlow) et un solveur industriel conventionnel aux Volumes Finis (OpenFOAM) sur le cas test du tourbillon de Taylor-Green, "
        "démontrant la conservation parfaite de l'énergie et l'absence totale de viscosité numérique parasite.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    p = doc.add_paragraph(
        "[Fig 2] représente un diagramme du front de Pareto illustrant le compromis entre la précision d'intégration temporelle "
        "(erreur résiduelle) et la latence de calcul par pas d'itération, comparant le présent procédé (intégrateurs exponentiels Lawson IF-RK2 "
        "et ETD-RK4) aux schémas classiques d'Euler et de Runge-Kutta conventionnels.",
        style="Inpi_ParagraphesNumérotés"
    )
    add_num_pr(p, num_id=2, ilvl=0)

    # 2. SECTION REVENDICATIONS
    doc.add_paragraph("Revendications", style="Inpi_Section")

    # Revendication 1 (Indépendante)
    p = doc.add_paragraph(
        "Procédé mis en œuvre par ordinateur pour le contrôle prédictif et l'asservissement en temps réel d'un écoulement fluide physique, "
        "comprenant les étapes suivantes exécutées en boucle par au moins un processeur :",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    p = doc.add_paragraph(
        "a) l'acquisition de signaux physiques issus de capteurs et leur transformation en un champ de vitesse dans un espace spectral de Fourier ;",
        style="Inpi_revendications"
    )
    p = doc.add_paragraph(
        "b) l'application itérative à chaque pas de temps d'un opérateur de dissipation spectrale modifié injectant une régularisation "
        "biharmonique d'ordre 4 aux hautes fréquences selon la relation D(k) = -ν |k|² [1 + α' |k|²], ledit opérateur constituant une borne "
        "matérielle stricte sur l'enstrophie empêchant tout dépassement de capacité des registres en virgule flottante ;",
        style="Inpi_revendications"
    )
    p = doc.add_paragraph(
        "c) l'application à chaque étape d'intégration d'une projection algébrique orthogonale de Leray-Helmholtz dans le domaine spectral, "
        "garantissant une divergence nulle stricte sans inversion de système matriciel creux en mémoire vive ;",
        style="Inpi_revendications"
    )
    p = doc.add_paragraph(
        "d) la conversion en temps réel du champ d'écoulement prédictif en un signal électrique de commande émis vers un actionneur physique "
        "régulant l'écoulement avec une latence inférieure à la milliseconde.",
        style="Inpi_revendications"
    )

    # Revendication 2
    p = doc.add_paragraph(
        "Procédé selon la revendication 1, caractérisé en ce que l'intégration temporelle numérique est effectuée par un opérateur d'intégration "
        "exponentielle analytique (ETD-RK4 ou IF-RK2 de Lawson) résolvant exactement la partie linéaire visqueuse, éliminant la restriction "
        "de pas de temps imposée par la condition de stabilité de Courant-Friedrichs-Lewy (CFL) visqueuse.",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    # Revendication 3
    p = doc.add_paragraph(
        "Procédé selon la revendication 1 ou 2, caractérisé en ce que le processeur calcule dynamiquement un indice de frustration triadique "
        "défini par le ratio entre la somme des modules des transferts d'énergie modaux et le module de la somme nette desdits transferts, "
        "l'atteinte d'un seuil prédéterminé dudit indice déclenchant une modification dynamique de la fréquence d'échantillonnage temporel "
        "ou de l'allocation des ressources de calcul du processeur.",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    # Revendication 4
    p = doc.add_paragraph(
        "Procédé selon l'une des revendications 1 à 3, caractérisé en ce que le processeur évalue en continu la préservation d'invariants physiques "
        "comprenant la circulation et l'hélicité, et déclenche une interruption matérielle de sécurité avec rejet de l'état calculé en cas de détection "
        "d'une discontinuité non physique, empêchant l'émission d'un signal de commande d'actionnement erroné.",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    # Revendication 5 (Dispositif)
    p = doc.add_paragraph(
        "Dispositif informatique de contrôle industriel embarqué pour équipement fluidique, caractérisé en ce qu'il comprend au moins "
        "un processeur microcontrôleur, une interface capteur connectée à l'équipement fluidique, une interface d'actionnement couplée "
        "à un actionneur physique dudit équipement, et une mémoire vive statique fonctionnant sans allocation dynamique de mémoire (mode no_std), "
        "ledit processeur exécutant le procédé selon l'une des revendications 1 à 4 pour asservir l'équipement fluidique avec une latence par "
        "itération inférieure à 100 microsecondes et une empreinte de mémoire vive inférieure à 3 Kilo-octets.",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    # Revendication 6 (Produit-programme)
    p = doc.add_paragraph(
        "Produit-programme d'ordinateur téléchargeable ou enregistré sur un support lisible par ordinateur, comprenant des instructions de code "
        "qui, lorsqu'elles sont exécutées par un processeur, conduisent celui-ci à mettre en œuvre les étapes du procédé selon l'une quelconque "
        "des revendications 1 à 4.",
        style="RevNumérotées"
    )
    add_num_pr(p, num_id=1, ilvl=0)

    # 3. SECTION ABRÉGÉ
    doc.add_paragraph("Abrégé", style="Inpi_Section")
    doc.add_paragraph(
        "L'invention concerne un procédé, un système informatique et un dispositif matériel pour la simulation et le contrôle en temps réel "
        "des équations de Navier-Stokes. Il met en œuvre une régularisation spectrale biharmonique fixant un plancher d'échelle ultraviolet "
        "qui empêche formellement tout plantage du processeur par singularité numérique, couplée à une projection d'incompressibilité algébrique "
        "directe et exacte (Leray-Helmholtz). Le procédé s'affranchit des résolutions matricielles creuses et des itérations de pression lentes "
        "de la CFD conventionnelle aux Volumes Finis. Présentant une empreinte mémoire vive inférieure à 3 Kilo-octets, une latence par itération "
        "inférieure à 100 microsecondes et un pare-feu mathématique anti-hallucination, l'invention est particulièrement destinée au contrôle "
        "cyber-physique embarqué d'actionneurs fluidiques sur microcontrôleurs industriels à ressources contraintes.",
        style="Inpi_Default"
    )
    # 4. PLANCHES DE DESSINS (FIGURES)
    if img_fig1 and os.path.exists(img_fig1):
        doc.add_paragraph("[Fig 1]", style="Inpi_BalisesConditionnelles")
        p_img1 = doc.add_paragraph(style="Inpi_ParagraphesNonNumérotés")
        run1 = p_img1.add_run()
        run1.add_picture(img_fig1, width=Inches(5.5))

    if img_fig2 and os.path.exists(img_fig2):
        doc.add_paragraph("[Fig 2]", style="Inpi_BalisesConditionnelles")
        p_img2 = doc.add_paragraph(style="Inpi_ParagraphesNonNumérotés")
        run2 = p_img2.add_run()
        run2.add_picture(img_fig2, width=Inches(5.5))

    doc.save(output_path)
    print(f"Document LibreOffice INPI généré avec succès : {output_path}")

def generate_msword_docx(template_path, output_path, math_dir, img_fig1=None, img_fig2=None):
    """Génère le document DOCX INPI au format Microsoft Word."""
    doc = docx.Document(template_path)
    clear_body(doc)

    # 1. SECTION DESCRIPTION
    doc.add_paragraph("Description", style="INPI_Titre_balise")
    
    # Titre de l'invention (strictement suivi du premier sous-titre)
    titre_text = ("Titre de l’invention : Procédé, système informatique et dispositif matériel "
                  "pour la simulation numérique et le contrôle prédictif en temps réel d'écoulements "
                  "fluides par régularisation spectrale à double échelle.")
    doc.add_paragraph(titre_text, style="INPI_Titre_balise")

    # Domaine technique
    add_subtitle(doc, "Domaine technique de l'invention", "INPI_Sous-titre")
    doc.add_paragraph(
        "La présente invention se rapporte au domaine de la mécanique des fluides numérique "
        "(CFD - Computational Fluid Dynamics), des systèmes cyber-physiques et de l'informatique embarquée. "
        "Plus particulièrement, l'invention concerne un procédé mis en œuvre par ordinateur pour la résolution "
        "stabilisée des équations de Navier-Stokes, optimisé pour s'exécuter sur des architectures matérielles "
        "à ressources limitées (microcontrôleurs « bare-metal ») afin de piloter et d'asservir des équipements "
        "physiques fluides en temps réel (ex : bioréacteurs, aérodynamique active, refroidissement thermique).",
        style="INPI_Paragraphe_numéroté"
    )

    # État de la technique
    add_subtitle(doc, "État de la technique", "INPI_Sous-titre")
    doc.add_paragraph(
        "La simulation d'écoulements turbulents pour le contrôle industriel repose sur les équations de Navier-Stokes. "
        "Le problème technique majeur réside dans la formation de structures tourbillonnaires à des échelles de plus en "
        "plus petites (phénomène d'étirement tourbillonnaire ou vortex stretching), ce qui conduit fréquemment dans les "
        "solveurs numériques à des singularités de calcul, des instabilités par accumulation d'énergie aux petites échelles, "
        "des dépassements de capacité produisant des erreurs fatales de type NaN (Not a Number), et par conséquent au plantage "
        "des systèmes informatiques de contrôle industriel.",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Pour tenter de contourner ces limitations, les logiciels de l'état de l'art (modèles LES, solveurs aux Volumes Finis "
        "tels qu'OpenFOAM ou ANSYS Fluent) nécessitent d'une part des résolutions itératives de systèmes matriciels creux "
        "(équation de Poisson pour la pression) exigeant plusieurs Gigaoctets de mémoire vive (RAM), ce qui interdit formellement "
        "tout déploiement embarqué sur microcontrôleurs pour un asservissement à faible latence (Edge Computing). D'autre part, "
        "ils recourent à des viscosités artificielles ou des filtres empiriques de sous-maille qui ne garantissent mathématiquement "
        "aucun bornage strict de l'enstrophie, laissant subsister un risque majeur d'interruption du processeur lors de fluctuations critiques.",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Il existe donc un besoin technique impérieux pour un procédé de calcul et de contrôle fluide garantissant informatiquement "
        "l'absence totale de dépassement de capacité (stabilité inconditionnelle et bornage de l'enstrophie), tout en réduisant drastiquement "
        "l'empreinte mémoire à quelques kilo-octets pour permettre un contrôle matériel en boucle fermée ultra-rapide.",
        style="INPI_Paragraphe_numéroté"
    )

    # Exposé de l'invention
    add_subtitle(doc, "Exposé de l'invention", "INPI_Sous-titre")
    doc.add_paragraph(
        "L'invention résout ce problème technique en introduisant un moteur de calcul pseudo-spectral couplé à une régularisation "
        "ultraviolette à plancher d'échelle (fondée sur une topologie de double échelle), implémenté sous la forme d'un code système "
        "déterministe sans allocation dynamique de mémoire (fonctionnant en mode no_std).",
        style="INPI_Paragraphe_numéroté"
    )

    # Description détaillée
    add_subtitle(doc, "Description détaillée", "INPI_Sous-titre")
    doc.add_paragraph(
        "Le procédé selon l'invention comprend les étapes techniques suivantes, exécutées en boucle temps réel par au moins un processeur :",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Étape A : Acquisition et Transformation spectrale. Le système acquiert les signaux de capteurs physiques (vitesses, pressions "
        "locales, consignes) et transforme le champ d'écoulement dans un espace spectral de Fourier à l'aide d'une transformée rapide (FFT).",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Étape B : Filtrage par Dissipation Modifiée (Plancher d'échelle). Le processeur applique itérativement un opérateur de dissipation "
        "spectrale D(k) configuré selon une loi modifiée à régularisation biharmonique d'ordre 4 aux hautes fréquences :",
        style="INPI_Paragraphe_numéroté"
    )

    # Balise mathématique 1 + Image de la formule
    doc.add_paragraph("[Math 1]", style="INPI_ParaNonNum")
    p_m1 = doc.add_paragraph(style="INPI_ParaNonNum")
    p_m1.add_run().add_picture(os.path.join(math_dir, "math1.png"), height=Inches(0.4))

    doc.add_paragraph(
        "où ν désigne la viscosité cinématique moléculaire, k le vecteur d'onde spatial, et α' un paramètre dimensionnel fixant un plancher "
        "d'échelle ultraviolet. L'effet technique direct de cette étape est la garantie informatique du maintien de l'enstrophie sous une borne "
        "stricte indépendante du maillage, interdisant mathématiquement et matériellement tout dépassement de capacité des registres en virgule flottante "
        "et tout plantage du processeur.",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Étape C : Intégration temporelle exacte et Projection de Leray-Helmholtz. Le processeur intègre temporellement l'état du fluide "
        "par un intégrateur exponentiel analytique (schéma ETD-RK4 ou IF-RK2 de Lawson) qui traite exactement la partie linéaire visqueuse, "
        "affranchissant le calcul de la condition de stabilité de Courant-Friedrichs-Lewy (CFL) visqueuse. À chaque pas, une projection algébrique "
        "orthogonale de Leray-Helmholtz est appliquée directement dans le domaine de Fourier selon :",
        style="INPI_Paragraphe_numéroté"
    )

    # Balise mathématique 2 + Image de la formule
    doc.add_paragraph("[Math 2]", style="INPI_ParaNonNum")
    p_m2 = doc.add_paragraph(style="INPI_ParaNonNum")
    p_m2.add_run().add_picture(os.path.join(math_dir, "math2.png"), height=Inches(0.4))

    doc.add_paragraph(
        "L'effet technique de cette projection algébrique directe est de garantir une divergence strictement nulle (incompressibilité à l'epsilon "
        "machine près, inférieure à 10^-16) pour chaque mode de Fourier, sans nécessiter le moindre stockage matriciel creux ni la moindre "
        "itération de solveur de Poisson en mémoire vive (supprimant les algorithmes itératifs de type SIMPLE ou PISO).",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Étape D : Capteur logiciel de Frustration Triadique. Le processeur calcule dynamiquement, à partir des transferts non linéaires d'énergie, "
        "un indice adimensionnel dit « Indice de Frustration Triadique », défini par le rapport :",
        style="INPI_Paragraphe_numéroté"
    )

    # Balise mathématique 3 + Image de la formule
    doc.add_paragraph("[Math 3]", style="INPI_ParaNonNum")
    p_m3 = doc.add_paragraph(style="INPI_ParaNonNum")
    p_m3.add_run().add_picture(os.path.join(math_dir, "math3.png"), height=Inches(0.4))

    doc.add_paragraph(
        "Cet indice mesure le degré de désalignement géométrique des triades de turbulence. Lorsque cet indice franchit un seuil prédéterminé, "
        "le système déclenche une allocation dynamique des ressources matérielles de calcul ou adapte la fréquence d'échantillonnage temporel.",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "Étape E : Asservissement matériel en boucle fermée. L'état prédit de l'écoulement est converti en temps réel en un signal électrique de commande "
        "transmis à un actionneur physique (e.g. variateur de vitesse d'agitateur de bioréacteur pour asservir le transfert massique kLa, volets "
        "d'aérodynamique active, pompes de refroidissement thermique).",
        style="INPI_Paragraphe_numéroté"
    )

    # Avantages techniques
    add_subtitle(doc, "Avantages techniques", "INPI_Sous-titre")
    p15_text = (
        "La synergie entre la régularisation spectrale biharmonique, la projection algébrique de Leray-Helmholtz et l'architecture sans allocation dynamique "
        "permet des performances matérielles inédites mesurées sur banc d'essai : "
        "(1) Une empreinte de mémoire vive (RAM) inférieure à 3 Kilo-octets (mesurée précisément à 2 624 octets pour une grille de dimensionnement embarqué), "
        "permettant une implantation directe dans la mémoire SRAM interne de microcontrôleurs basse consommation (architectures ARM Cortex-M4/M7 ou RISC-V), "
        "sans mémoire externe ; "
        "(2) Un déterminisme d'exécution temps réel strict avec une latence médiane par itération inférieure à 100 microsecondes (mesurée à 59,8 µs), "
        "autorisant des boucles de contrôle cyber-physiques à plus de 10 000 Hz."
    )
    doc.add_paragraph(p15_text, style="INPI_Paragraphe_numéroté")

    # Optimisation autonome et benchmarks
    add_subtitle(doc, "Validation empirique certifiée", "INPI_Sous-titre")
    doc.add_paragraph(
        "Le procédé s'intègre en outre dans un système d'optimisation par recherche autonome et modèles d'ordre réduit (ROMs), "
        "où chaque configuration paramétrique est validée vis-à-vis d'invariants formels certifiés par démonstrateur interactif de théorèmes (Lean 4). "
        "L'ensemble des résultats empiriques a été scellé par une empreinte cryptographique SHA-256 (Sceau : 0ae0e5d97da424e0).",
        style="INPI_Paragraphe_numéroté"
    )

    p17_text = (
        "Les gains quantitatifs établis par rapport aux solveurs conventionnels (Volumes Finis type OpenFOAM) comprennent : "
        "(i) Précision numérique absolue (Zéro diffusion numérique) : Sur le cas test standardisé du Tourbillon de Taylor-Green (UC7), "
        "l'erreur L2 sur la vitesse est mesurée à 7,24 x 10^-14 sur grille 128², prouvant l'éradication totale de la viscosité artificielle de discrétisation ; "
        "(ii) Vitesse algorithmique (Contournement de la limite CFL) : Résolution de la cavité entraînée (UC8) en 0,59 seconde et de la cascade "
        "de Kolmogorov en 3D (UC11) en 0,21 seconde grâce au schéma intégrateur exponentiel ETD-RK4, soit un gain de vitesse de calcul supérieur à 10× "
        "par rapport aux algorithmes itératifs de pression ; "
        "(iii) Sécurité épistémique matérielle (Pare-feu mathématique anti-hallucination) : Soumis à des conditions aux limites non physiques "
        "(e.g. UC15 - tore périodique avec vorticité moyenne non nulle forçant un paradoxe mathématique), le procédé rejette formellement l'état "
        "en mesurant une perte de circulation de ~100% et déclenche une interruption de sécurité matérielle « Échec Contrôle Négatif », "
        "empêchant toute prise de commande erronée par l'actionneur physique."
    )
    doc.add_paragraph(p17_text, style="INPI_Paragraphe_numéroté")

    # Brève description des dessins
    add_subtitle(doc, "Brève description des dessins", "INPI_Sous-titre")
    doc.add_paragraph(
        "[Fig 1] représente un graphique comparatif de la décroissance temporelle de l'énergie cinétique entre le procédé selon l'invention "
        "(LeanFlow) et un solveur industriel conventionnel aux Volumes Finis (OpenFOAM) sur le cas test du tourbillon de Taylor-Green, "
        "démontrant la conservation parfaite de l'énergie et l'absence totale de viscosité numérique parasite.",
        style="INPI_Paragraphe_numéroté"
    )

    doc.add_paragraph(
        "[Fig 2] représente un diagramme du front de Pareto illustrant le compromis entre la précision d'intégration temporelle "
        "(erreur résiduelle) et la latence de calcul par pas d'itération, comparant le présent procédé (intégrateurs exponentiels Lawson IF-RK2 "
        "et ETD-RK4) aux schémas classiques d'Euler et de Runge-Kutta conventionnels.",
        style="INPI_Paragraphe_numéroté"
    )

    # 2. SECTION REVENDICATIONS
    doc.add_paragraph("Revendications", style="INPI_Titre_balise")

    # Revendication 1
    doc.add_paragraph(
        "Procédé mis en œuvre par ordinateur pour le contrôle prédictif et l'asservissement en temps réel d'un écoulement fluide physique, "
        "comprenant les étapes suivantes exécutées en boucle par au moins un processeur :",
        style="INPI_Revendications"
    )

    doc.add_paragraph(
        "a) l'acquisition de signaux physiques issus de capteurs et leur transformation en un champ de vitesse dans un espace spectral de Fourier ;",
        style="INPI_RevNonNum"
    )
    doc.add_paragraph(
        "b) l'application itérative à chaque pas de temps d'un opérateur de dissipation spectrale modifié injectant une régularisation "
        "biharmonique d'ordre 4 aux hautes fréquences selon la relation D(k) = -ν |k|² [1 + α' |k|²], ledit opérateur constituant une borne "
        "matérielle stricte sur l'enstrophie empêchant tout dépassement de capacité des registres en virgule flottante ;",
        style="INPI_RevNonNum"
    )
    doc.add_paragraph(
        "c) l'application à chaque étape d'intégration d'une projection algébrique orthogonale de Leray-Helmholtz dans le domaine spectral, "
        "garantissant une divergence nulle stricte sans inversion de système matriciel creux en mémoire vive ;",
        style="INPI_RevNonNum"
    )
    doc.add_paragraph(
        "d) la conversion en temps réel du champ d'écoulement prédictif en un signal électrique de commande émis vers un actionneur physique "
        "régulant l'écoulement avec une latence inférieure à la milliseconde.",
        style="INPI_RevNonNum"
    )

    # Revendication 2
    doc.add_paragraph(
        "Procédé selon la revendication 1, caractérisé en ce que l'intégration temporelle numérique est effectuée par un opérateur d'intégration "
        "exponentielle analytique (ETD-RK4 ou IF-RK2 de Lawson) résolvant exactement la partie linéaire visqueuse, éliminant la restriction "
        "de pas de temps imposée par la condition de stabilité de Courant-Friedrichs-Lewy (CFL) visqueuse.",
        style="INPI_Revendications"
    )

    # Revendication 3
    doc.add_paragraph(
        "Procédé selon la revendication 1 ou 2, caractérisé en ce que le processeur calcule dynamiquement un indice de frustration triadique "
        "défini par le ratio entre la somme des modules des transferts d'énergie modaux et le module de la somme nette desdits transferts, "
        "l'atteinte d'un seuil prédéterminé dudit indice déclenchant une modification dynamique de la fréquence d'échantillonnage temporel "
        "ou de l'allocation des ressources de calcul du processeur.",
        style="INPI_Revendications"
    )

    # Revendication 4
    doc.add_paragraph(
        "Procédé selon l'une des revendications 1 à 3, caractérisé en ce que le processeur évalue en continu la préservation d'invariants physiques "
        "comprenant la circulation et l'hélicité, et déclenche une interruption matérielle de sécurité avec rejet de l'état calculé en cas de détection "
        "d'une discontinuité non physique, empêchant l'émission d'un signal de commande d'actionnement erroné.",
        style="INPI_Revendications"
    )

    # Revendication 5
    doc.add_paragraph(
        "Dispositif informatique de contrôle industriel embarqué pour équipement fluidique, caractérisé en ce qu'il comprend au moins "
        "un processeur microcontrôleur, une interface capteur connectée à l'équipement fluidique, une interface d'actionnement couplée "
        "à un actionneur physique dudit équipement, et une mémoire vive statique fonctionnant sans allocation dynamique de mémoire (mode no_std), "
        "ledit processeur exécutant le procédé selon l'une des revendications 1 à 4 pour asservir l'équipement fluidique avec une latence par "
        "itération inférieure à 100 microsecondes et une empreinte de mémoire vive inférieure à 3 Kilo-octets.",
        style="INPI_Revendications"
    )

    # Revendication 6
    doc.add_paragraph(
        "Produit-programme d'ordinateur téléchargeable ou enregistré sur un support lisible par ordinateur, comprenant des instructions de code "
        "qui, lorsqu'elles sont exécutées par un processeur, conduisent celui-ci à mettre en œuvre les étapes du procédé selon l'une quelconque "
        "des revendications 1 à 4.",
        style="INPI_Revendications"
    )

    # 3. SECTION ABRÉGÉ
    doc.add_paragraph("Abrégé", style="INPI_Titre_balise")
    doc.add_paragraph(
        "L'invention concerne un procédé, un système informatique et un dispositif matériel pour la simulation et le contrôle en temps réel "
        "des équations de Navier-Stokes. Il met en œuvre une régularisation spectrale biharmonique fixant un plancher d'échelle ultraviolet "
        "qui empêche formellement tout plantage du processeur par singularité numérique, couplée à une projection d'incompressibilité algébrique "
        "directe et exacte (Leray-Helmholtz). Le procédé s'affranchit des résolutions matricielles creuses et des itérations de pression lentes "
        "de la CFD conventionnelle aux Volumes Finis. Présentant une empreinte mémoire vive inférieure à 3 Kilo-octets, une latence par itération "
        "inférieure à 100 microsecondes et un pare-feu mathématique anti-hallucination, l'invention est particulièrement destinée au contrôle "
        "cyber-physique embarqué d'actionneurs fluidiques sur microcontrôleurs industriels à ressources contraintes.",
        style="Text body"
    )

    # 4. PLANCHES DE DESSINS (FIGURES)
    if img_fig1 and os.path.exists(img_fig1):
        doc.add_paragraph("[Fig 1]", style="INPI_ParaNonNum")
        p_img1 = doc.add_paragraph(style="INPI_ParaNonNum")
        run1 = p_img1.add_run()
        run1.add_picture(img_fig1, width=Inches(5.5))

    if img_fig2 and os.path.exists(img_fig2):
        doc.add_paragraph("[Fig 2]", style="INPI_ParaNonNum")
        p_img2 = doc.add_paragraph(style="INPI_ParaNonNum")
        run2 = p_img2.add_run()
        run2.add_picture(img_fig2, width=Inches(5.5))

    doc.save(output_path)
    print(f"Document MS Word INPI généré avec succès : {output_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    math_dir = os.path.join(base_dir, "figures")
    
    # 1. Génération des images de formules mathématiques
    ensure_math_images(math_dir)

    tpl_lo = os.path.join(base_dir, "template", "Exemple_brochure_LibreOffice.docx")
    tpl_ms = os.path.join(base_dir, "template", "Exemple_brochure_MSWord.docx")
    
    fig1_path = os.path.join(os.path.dirname(base_dir), "marketing", "energy_decay.png")
    fig2_path = os.path.join(os.path.dirname(base_dir), "marketing", "pareto_front.png")

    # Chemins de sortie
    out_lo = os.path.join(base_dir, "demande_brevet_INPI.docx")
    out_lo_alt = os.path.join(base_dir, "demande_brevet.docx")
    out_ms = os.path.join(base_dir, "demande_brevet_INPI_MSWord.docx")

    print("Génération du document INPI (LibreOffice)...")
    generate_libreoffice_docx(tpl_lo, out_lo, math_dir, fig1_path, fig2_path)

    print("Génération du document INPI (Microsoft Word)...")
    generate_msword_docx(tpl_ms, out_ms, math_dir, fig1_path, fig2_path)
    shutil.copy(out_ms, out_lo_alt)

    print("Tous les documents INPI ont été régénérés avec succès.")
